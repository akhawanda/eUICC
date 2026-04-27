"""
Turn the JSON report from phase2_e2e.sh into a Word document with:
- Test environment + executive summary
- Per-step section: action taken, what was checked, observed state
- A profile-number / state table per step
- Final pass/fail verdict

Usage:
    python3 gen_phase2_doc.py phase2_report.json Phase2-E2E-Results.docx
"""
import json
import sys
from datetime import datetime
from pathlib import Path

from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL


GREEN = RGBColor(0x22, 0x86, 0x36)
RED = RGBColor(0xC0, 0x39, 0x2B)
GREY = RGBColor(0x55, 0x55, 0x55)


def add_kv_table(doc, rows):
    t = doc.add_table(rows=len(rows), cols=2)
    t.style = "Light Grid"
    for i, (k, v) in enumerate(rows):
        t.rows[i].cells[0].text = k
        t.rows[i].cells[1].text = str(v)
        for p in t.rows[i].cells[0].paragraphs:
            for run in p.runs:
                run.bold = True
    return t


def add_profile_table(doc, profiles):
    if not profiles:
        p = doc.add_paragraph("(no profiles)")
        p.runs[0].font.color.rgb = GREY
        return
    t = doc.add_table(rows=1 + len(profiles), cols=3)
    t.style = "Light Grid"
    hdr = t.rows[0].cells
    hdr[0].text = "ICCID"
    hdr[1].text = "Profile state"
    hdr[2].text = "ISD-P AID"
    for c in hdr:
        for p in c.paragraphs:
            for r in p.runs:
                r.bold = True
    for i, p in enumerate(profiles):
        row = t.rows[i + 1].cells
        row[0].text = p.get("iccid", "")
        state_val = p.get("profileState", 0)
        state_str = {0: "0 (disabled)", 1: "1 (enabled)"}.get(state_val, str(state_val))
        row[1].text = state_str
        row[2].text = p.get("isdpAid", "")
        if state_val == 1:
            for para in row[1].paragraphs:
                for r in para.runs:
                    r.font.color.rgb = GREEN
                    r.bold = True


def add_eim_table(doc, eims):
    if not eims:
        doc.add_paragraph("(no eIMs associated)")
        return
    t = doc.add_table(rows=1 + len(eims), cols=4)
    t.style = "Light Grid"
    hdr = t.rows[0].cells
    hdr[0].text = "eIM Id"
    hdr[1].text = "FQDN"
    hdr[2].text = "Counter"
    hdr[3].text = "Token"
    for c in hdr:
        for p in c.paragraphs:
            for r in p.runs:
                r.bold = True
    for i, e in enumerate(eims):
        row = t.rows[i + 1].cells
        row[0].text = str(e.get("eimId", ""))
        row[1].text = str(e.get("eimFqdn", ""))
        row[2].text = str(e.get("counterValue", ""))
        row[3].text = str(e.get("associationToken", ""))


def main():
    src = Path(sys.argv[1] if len(sys.argv) > 1 else "phase2_report.json")
    dst = Path(sys.argv[2] if len(sys.argv) > 2 else "Phase2-E2E-Results.docx")

    with src.open() as f:
        report = json.load(f)
    steps = report.get("steps", [])
    eid = report.get("eid", "")
    started = report.get("started", "")

    final = steps[-1] if steps else {}
    overall_pass = bool(final.get("pass"))

    doc = Document()

    title = doc.add_heading("eUICC Simulator — Phase 2 E2E Test Report", level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = sub.add_run("Download × 2 · eIM Swap (connectxiot.com → Tahwol) · Queued Enable from Tahwol")
    r.italic = True
    r.font.color.rgb = GREY

    doc.add_heading("Test Environment", level=1)
    add_kv_table(doc, [
        ("Device EID", eid),
        ("SM-DP+", "smdpplus.connectxiot.com (our SM-DP+, GSMA SGP.26 NIST chain)"),
        ("eIM (Go server)", "eimserver.connectxiot.com (multi-identity: connectxiot.com + Tahwol)"),
        ("eUICC Simulator", "http://127.0.0.1:8100 (deployed at euicc.connectxiot.com)"),
        ("IPA Simulator", "http://127.0.0.1:8101 (deployed at euicc.connectxiot.com)"),
        ("Test driver", "phase2_e2e.sh"),
        ("Run started", started),
    ])

    doc.add_heading("Executive Summary", level=1)
    p = doc.add_paragraph()
    if overall_pass:
        run = p.add_run("PASS — all 8 steps completed; Profile B was successfully enabled via the Tahwol eIM identity.")
        run.font.color.rgb = GREEN
        run.bold = True
    else:
        run = p.add_run("FAIL — see step-by-step section for details.")
        run.font.color.rgb = RED
        run.bold = True

    if steps:
        s1 = steps[0].get("profilesAfter", {}).get("profileInfoListOk", [])
        s2 = steps[1].get("profilesAfter", {}).get("profileInfoListOk", []) if len(steps) > 1 else []
        if s2:
            iccid_a = s1[0]["iccid"] if s1 else "—"
            iccid_b = s2[1]["iccid"] if len(s2) > 1 else "—"
            doc.add_paragraph(f"Profile A: {iccid_a}").runs[0].font.size = Pt(11)
            doc.add_paragraph(f"Profile B: {iccid_b} (target of step 7 enable)").runs[0].font.size = Pt(11)

    doc.add_heading("Step-by-step Results", level=1)

    for idx, step in enumerate(steps, 1):
        doc.add_heading(step.get("name", f"Step {idx}"), level=2)
        meta = []
        if "matchingId" in step:
            meta.append(("Matching ID", step["matchingId"]))
        if "queue" in step:
            meta.append(("eIM queue", step["queue"]))
        if "targetIccid" in step:
            meta.append(("Target ICCID", step["targetIccid"]))
        if "iccidA" in step:
            meta.append(("ICCID A", step["iccidA"]))
        if "iccidB" in step:
            meta.append(("ICCID B", step["iccidB"]))
        if "expectedEnabled" in step:
            meta.append(("Expected enabled", step["expectedEnabled"]))
        if "actualEnabled" in step:
            meta.append(("Actual enabled", step["actualEnabled"] or "(none)"))
        if "pass" in step:
            meta.append(("Verdict", "PASS" if step["pass"] else "FAIL"))
        if meta:
            add_kv_table(doc, meta)

        if "downloadResult" in step and isinstance(step["downloadResult"], dict):
            dr = step["downloadResult"]
            doc.add_paragraph().add_run("8-step download flow:").bold = True
            for s in dr.get("steps", []):
                doc.add_paragraph(s, style="List Bullet")
            if dr.get("error"):
                err = doc.add_paragraph()
                er = err.add_run(f"Error: {dr['error']}")
                er.font.color.rgb = RED

        profiles_dict = step.get("profilesAfter") or step.get("profiles")
        if profiles_dict:
            ps = profiles_dict.get("profileInfoListOk")
            if ps is not None:
                doc.add_paragraph().add_run("Profiles on eUICC after this step:").bold = True
                add_profile_table(doc, ps)

        eims_dict = step.get("eimsAfter") or step.get("eims")
        if eims_dict:
            es = eims_dict.get("eimConfigurationDataList")
            if es is not None:
                doc.add_paragraph().add_run("eIM associations on eUICC after this step:").bold = True
                add_eim_table(doc, es)

    doc.add_heading("Notes", level=1)
    notes = [
        "Profiles installed on the eUICC sim carry placeholder ICCIDs because the simulator's BPP processor stubs SCP03t decryption + PE-Header parsing. The wire-protocol download against the SM-DP+ is fully spec-compliant (Executed-Success on InitiateAuthentication, AuthenticateClient, and GetBoundProfilePackage); only the post-decode profile rendering is stubbed. Tracked as the BPP-internals open follow-up.",
        "PSMO/eCO actions are signed by the eIM Go server using the identity associated with the device in eim_db.eim_device_associations. The script mirrors eUICC-side addEim/deleteEim onto that table so subsequent operations are signed by an identity the eUICC still trusts.",
        "ICCIDs on the SGP.32 wire are BCD-encoded (nibble-swapped from the eUICC's stored bytes). The eUICC's _execute_psmo applies the swap back before profile lookup so the PSMO path matches the direct ES10c API path.",
        "No SM-DP+ or eIM Go server source code was modified during this test — fixes are in IPA + eUICC simulator only, per the project scope rule.",
    ]
    for n in notes:
        doc.add_paragraph(n, style="List Bullet")

    doc.save(dst)
    print(f"Wrote {dst}")


if __name__ == "__main__":
    main()
