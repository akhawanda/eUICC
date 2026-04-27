"""Generate "RSP Certificates & Authentication Guide.docx" with embedded diagrams.

Run on server:
  python3 /tmp/build_cert_guide.py /tmp/cert_guide.docx
"""

import sys, os
import matplotlib
matplotlib.use("Agg")
import matplotlib.pyplot as plt
import matplotlib.patches as mpatches
from matplotlib.patches import FancyBboxPatch, FancyArrowPatch
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


OUT_DIR = os.path.dirname(os.path.abspath(sys.argv[1]))
IMG_DIR = os.path.join(OUT_DIR, "_diagrams")
os.makedirs(IMG_DIR, exist_ok=True)


# ===========================================================================
# DIAGRAMS
# ===========================================================================


def fig_trust_hierarchy(path: str) -> None:
    fig, ax = plt.subplots(figsize=(11, 7), dpi=140)
    ax.set_xlim(0, 12); ax.set_ylim(0, 8); ax.axis("off")

    def box(x, y, w, h, label, color, fontcolor="black", fontweight="normal", fontsize=10):
        rect = FancyBboxPatch((x, y), w, h, boxstyle="round,pad=0.05,rounding_size=0.15",
                              fc=color, ec="#333", lw=1.4)
        ax.add_patch(rect)
        ax.text(x + w/2, y + h/2, label, ha="center", va="center",
                color=fontcolor, fontsize=fontsize, fontweight=fontweight, wrap=True)

    # Title
    ax.text(6, 7.6, "RSP Trust Hierarchy", ha="center", fontsize=14, fontweight="bold")
    ax.text(6, 7.25, "(GSMA SGP.22 v3 / SGP.26 NIST P-256 Test profile)", ha="center", fontsize=9, color="#555")

    # GSMA Test CI (top)
    box(4.0, 5.6, 4.0, 1.2,
        "GSMA SGP.26 Test CI (NIST P-256)\nCN=Test CI, OU=TESTCERT, O=RSPTEST\nSKI = F5417...85C3",
        "#FFF1A8", fontweight="bold", fontsize=9)

    # ConnectX self-signed CI (top right, alternate trust root)
    box(8.5, 5.6, 3.3, 1.2,
        "ConnectX self-signed CI\n(test-only, no GSMA chain)\nSKI = CF26C...FE80",
        "#E8E8E8", fontsize=9)

    # SM-DP+ side certs
    box(0.4, 3.6, 2.7, 1.0,
        "CERT.DPauth.SIG\n(SM-DP+ identity)\nfor ServerSigned1",
        "#CDE7FF", fontsize=9)
    box(3.4, 3.6, 2.7, 1.0,
        "CERT.DPpb.SIG\n(profile binding)\nfor SmdpSigned2",
        "#CDE7FF", fontsize=9)

    # eUICC side
    box(7.4, 3.6, 2.4, 1.0,
        "EUM cert\n(eUICC manufacturer)",
        "#D4F4D4", fontsize=9)

    box(8.0, 1.6, 2.4, 1.0,
        "eUICC cert\n(per-EID, end-entity)\nfor EuiccSigned1",
        "#D4F4D4", fontsize=9)

    # Arrows: GSMA CI signs SM-DP+ certs
    ax.annotate("", xy=(1.7, 4.6), xytext=(5.0, 5.6),
                arrowprops=dict(arrowstyle="->", color="#333", lw=1.4))
    ax.annotate("", xy=(4.7, 4.6), xytext=(5.6, 5.6),
                arrowprops=dict(arrowstyle="->", color="#333", lw=1.4))
    ax.text(2.7, 5.05, "issued by\nGSMA Test CI", fontsize=8, color="#555")

    # Arrows: ConnectX CI signs EUM
    ax.annotate("", xy=(8.6, 4.6), xytext=(9.7, 5.6),
                arrowprops=dict(arrowstyle="->", color="#333", lw=1.4))
    ax.text(9.0, 5.05, "issued by\nConnectX CI", fontsize=8, color="#555")

    # Arrow: EUM signs eUICC
    ax.annotate("", xy=(9.2, 2.6), xytext=(8.6, 3.6),
                arrowprops=dict(arrowstyle="->", color="#333", lw=1.4))
    ax.text(8.7, 3.0, "signs", fontsize=8, color="#555")

    # eUICC trusts BOTH roots dashed line
    ax.annotate("", xy=(6.0, 6.2), xytext=(8.5, 6.2),
                arrowprops=dict(arrowstyle="<->", color="#a00", lw=1.0, ls="--"))
    ax.text(7.25, 6.45, "eUICC trusts BOTH roots\n(verification list)", fontsize=8, color="#a00",
            ha="center", style="italic")

    # Legend
    ax.text(0.5, 0.6, "Solid arrow = X.509 cert signing relationship    Dashed = trust-list membership",
            fontsize=8, color="#333")
    ax.text(0.5, 0.25, "Color: yellow=GSMA root, gray=local root, blue=SM-DP+, green=eUICC",
            fontsize=8, color="#333")

    plt.tight_layout()
    plt.savefig(path, bbox_inches="tight"); plt.close(fig)


def fig_certs_per_actor(path: str) -> None:
    fig, ax = plt.subplots(figsize=(11, 6), dpi=140)
    ax.set_xlim(0, 12); ax.set_ylim(0, 7); ax.axis("off")
    ax.text(6, 6.5, "Certificate Material per Actor", ha="center", fontsize=14, fontweight="bold")

    # 3 columns — eIM, SM-DP+, eUICC
    cols = [
        (0.3, 4.0, "eIM\n(eimserver.connectxiot.com)", "#FCE7B5", [
            "• eIM signing keypair (private_key, public_key)",
            "  — used to sign EuiccPackageRequest (BF51 → A0)",
            "  — registered in DB as eim_registry row",
            "• Per-eIM identity advertised to eUICC",
            "  via AddEim / EuiccPackageResult eimId field",
            "• NO X.509 chain involvement on the wire",
            "• No GSMA TestCI dependency",
        ]),
        (4.2, 4.0, "SM-DP+\n(smdpplus.connectxiot.com)", "#CDE7FF", [
            "• Trust anchor: GSMA SGP.26 NIST P-256 Test CI",
            "  ci_nist.crt (PUBLIC only)",
            "• CERT.DPauth.SIG + private key — signs",
            "  ServerSigned1 in InitiateAuthentication",
            "• CERT.DPpb.SIG + private key — signs",
            "  SmdpSigned2 in AuthenticateClient",
            "  (binds the eUICC to the BPP)",
            "• TLS cert (Let's Encrypt) for HTTPS only",
        ]),
        (8.1, 4.0, "eUICC\n(per EID, in sim memory)", "#D4F4D4", [
            "• Trust anchors (verification): TWO roots",
            "  1) GSMA SGP.26 Test CI public cert",
            "     (so SM-DP+ chain is recognised)",
            "  2) ConnectX self-signed CI",
            "     (sim's own root)",
            "• Signing chain (own private key only):",
            "  ConnectX-CI → ConnectX-EUM → eUICC",
            "• eUICC private key signs EuiccSigned1/2",
        ]),
    ]
    for x, y, title, color, lines in cols:
        box = FancyBboxPatch((x, 0.3), 3.7, 5.5, boxstyle="round,pad=0.05,rounding_size=0.2",
                             fc=color, ec="#333", lw=1.4)
        ax.add_patch(box)
        ax.text(x + 1.85, 5.45, title, ha="center", fontsize=11, fontweight="bold")
        for i, line in enumerate(lines):
            ax.text(x + 0.15, 5.0 - i * 0.45, line, fontsize=8.5, va="top")

    plt.tight_layout()
    plt.savefig(path, bbox_inches="tight"); plt.close(fig)


def fig_es10b_auth_sequence(path: str) -> None:
    fig, ax = plt.subplots(figsize=(13, 11), dpi=140)
    ax.set_xlim(0, 13); ax.set_ylim(0, 14); ax.axis("off")
    ax.text(6.5, 13.5, "Profile Download — 8-step ES9+/ES10b Authentication Flow",
            ha="center", fontsize=14, fontweight="bold")
    ax.text(6.5, 13.1, "(GSMA SGP.22 v3 §5.6 + §5.7)",
            ha="center", fontsize=9, color="#555")

    # Lanes
    lanes = [(1.5, "eIM"), (4.0, "IPA"), (7.0, "eUICC"), (10.5, "SM-DP+")]
    for x, label in lanes:
        rect = FancyBboxPatch((x - 0.85, 12.0), 1.7, 0.55, boxstyle="round,pad=0.04,rounding_size=0.1",
                              fc="#333", ec="none")
        ax.add_patch(rect)
        ax.text(x, 12.27, label, ha="center", va="center", color="white", fontsize=10, fontweight="bold")
        ax.plot([x, x], [0.5, 12.0], color="#bbb", ls="--", lw=0.8, zorder=0)

    def msg(y, x_from, x_to, text, color="#333", note=None, dashed=False):
        ax.annotate("", xy=(x_to, y), xytext=(x_from, y),
                    arrowprops=dict(arrowstyle="->", color=color, lw=1.4,
                                    linestyle="--" if dashed else "-"))
        x_mid = (x_from + x_to) / 2
        ax.text(x_mid, y + 0.1, text, ha="center", fontsize=8.5, color=color, fontweight="bold")
        if note:
            ax.text(x_mid, y - 0.20, note, ha="center", fontsize=7.5, color="#555", style="italic")

    def section(y, text):
        ax.text(0.2, y, text, fontsize=9, fontweight="bold", color="#1a4ea0")

    # Setup
    section(11.6, "Phase 1 — eIM triggers download (SGP.32 ESipa)")
    msg(11.2, 4.0, 1.5, "1. ESipa.GetEimPackage (poll)")
    msg(10.7, 1.5, 4.0, "BF54 ProfileDownloadTrigger",
        note="activationCode = '1$smdpplus.../$<matchingId>'", dashed=True)

    section(10.1, "Phase 2 — eUICC challenge & info (ES10b)")
    msg(9.6, 4.0, 7.0, "2. ES10b.GetEuiccChallenge")
    msg(9.1, 7.0, 4.0, "16B random Nonce", dashed=True)
    msg(8.6, 4.0, 7.0, "3. ES10b.GetEuiccInfo1")
    msg(8.1, 7.0, 4.0, "BF20 EUICCInfo1 + euiccCiPKIdListForVerification",
        note="lists trusted CI roots: GSMA TestCI + ConnectX CI", dashed=True)

    section(7.6, "Phase 3 — Server side authentication (ES9+)")
    msg(7.1, 4.0, 10.5, "4. ES9+.InitiateAuthentication",
        note="JSON: euiccChallenge (base64), euiccInfo1, smdpAddress")
    msg(6.6, 10.5, 4.0,
        "transactionId, ServerSigned1, sig1 (DPauth), DPauth cert, ciPKID",
        note="ServerSigned1 = [80]txId [81]euiccChall [83]srvAddr [84]srvChall  (SGP.22 v3 layout)",
        dashed=True)

    msg(6.0, 4.0, 7.0, "5. ES10b.AuthenticateServer",
        note="forwards ServerSigned1 + sig + DPauth cert + ctxParams1(matchingId)")
    ax.text(7.5, 5.6, "  eUICC verifies:", fontsize=8, color="#a00")
    ax.text(7.6, 5.4, "    • ciPKID in trusted list", fontsize=7.5, color="#a00")
    ax.text(7.6, 5.22, "    • cert chain to CI root", fontsize=7.5, color="#a00")
    ax.text(7.6, 5.04, "    • ECDSA(serverPubKey, sig1, ServerSigned1Raw)", fontsize=7.5, color="#a00")
    ax.text(7.6, 4.86, "    • euiccChallenge echoed", fontsize=7.5, color="#a00")

    msg(4.6, 7.0, 4.0, "EuiccSigned1, sig1 (eUICC), eUICC cert, EUM cert",
        note="EuiccSigned1 = [80]txId [83]srvAddr [84]srvChall [BF22]info2 ctxParams1",
        dashed=True)

    section(4.1, "Phase 4 — Client authentication & key exchange (ES9+)")
    msg(3.6, 4.0, 10.5, "6. ES9+.AuthenticateClient",
        note="IPA wraps as BF38→A0→{EuiccSigned1Raw + 5F37 sig + euiccCert + eumCert}")
    ax.text(10.7, 3.0, "SM-DP+ verifies:", fontsize=8, color="#a00")
    ax.text(10.8, 2.83, "  • EUM chain to GSMA CI", fontsize=7.5, color="#a00")
    ax.text(10.8, 2.66, "  • eUICC chain to EUM", fontsize=7.5, color="#a00")
    ax.text(10.8, 2.49, "  • sig over EuiccSigned1", fontsize=7.5, color="#a00")
    ax.text(10.8, 2.32, "  • matchingId → released profile", fontsize=7.5, color="#a00")
    msg(2.0, 10.5, 4.0,
        "SmdpSigned2 + sig2 (DPpb) + DPpb cert + profileMetadata",
        note="sig2 over SmdpSigned2 || 5F37 40 || euiccSig1   (binds the BPP to the eUICC)",
        dashed=True)

    section(1.5, "Phase 5 — Bound profile install (ES9+ → ES10b)")
    msg(1.0, 4.0, 7.0, "7. ES10b.PrepareDownload + 8. LoadBoundProfilePackage",
        note="eUICC verifies sig2 (DPpb cert), derives SCP03t keys, installs PE-Header → ProfileInstallationResult")

    plt.tight_layout()
    plt.savefig(path, bbox_inches="tight"); plt.close(fig)


def fig_signature_layout(path: str) -> None:
    fig, ax = plt.subplots(figsize=(11, 6.5), dpi=140)
    ax.set_xlim(0, 12); ax.set_ylim(0, 7); ax.axis("off")
    ax.text(6, 6.6, "What gets signed by what — TBS layouts", ha="center", fontsize=14, fontweight="bold")

    def block(y, label, components, color, comp_color="#fff"):
        ax.text(0.3, y + 0.45, label, fontsize=10, fontweight="bold")
        x = 0.3
        for w, txt in components:
            rect = FancyBboxPatch((x, y - 0.05), w, 0.6,
                                  boxstyle="round,pad=0.03,rounding_size=0.06",
                                  fc=color, ec="#333", lw=1.0)
            ax.add_patch(rect)
            ax.text(x + w/2, y + 0.25, txt, ha="center", va="center", fontsize=8)
            x += w + 0.05

    block(5.4, "Server: ServerSigned1 (signed by CERT.DPauth.SIG)",
          [(2.2, "[80] transactionId"),
           (2.4, "[81] euiccChallenge"),
           (3.4, "[83] serverAddress (UTF8)"),
           (2.4, "[84] serverChallenge")],
          "#FFE4A8")
    ax.text(0.3, 5.05, "TBS = the SEQUENCE bytes above (verbatim from SM-DP+ wire — never re-encoded by eUICC)",
            fontsize=8, color="#555", style="italic")

    block(3.6, "eUICC: EuiccSigned1 (signed by eUICC private key)",
          [(2.2, "[80] transactionId"),
           (3.4, "[83] serverAddress (UTF8)"),
           (2.4, "[84] serverChallenge"),
           (1.7, "[BF22] info2"),
           (1.5, "ctxParams1")],
          "#D4F4D4")

    block(1.6, "Server: SmdpSigned2 + euiccSig1  (signed by CERT.DPpb.SIG)",
          [(2.2, "[80] transactionId"),
           (2.0, "ccRequiredFlag (BOOL)"),
           (2.6, "[5F49] bppEuiccOtpk OPT"),
           (1.0, "‖"),
           (2.5, "5F 37 40 + euiccSig1 (67B)")],
          "#FFE4A8")
    ax.text(0.3, 1.25, "TBS = SmdpSigned2_DER  ||  5F37 40  ||  euiccSig1   ← SGP.22 §5.6.2 binds SM-DP+ response to eUICC's prior signature",
            fontsize=8, color="#a00", style="italic")

    plt.tight_layout()
    plt.savefig(path, bbox_inches="tight"); plt.close(fig)


def fig_eim_relay(path: str) -> None:
    fig, ax = plt.subplots(figsize=(11, 6), dpi=140)
    ax.set_xlim(0, 12); ax.set_ylim(0, 7); ax.axis("off")
    ax.text(6, 6.5, "ESipa Relay (PSMO/eCO ops via eIM)",
            ha="center", fontsize=14, fontweight="bold")
    ax.text(6, 6.15, "(GSMA SGP.32 — eIM signs once, eUICC verifies counter)",
            ha="center", fontsize=9, color="#555")

    lanes = [(1.5, "eIM"), (5.0, "IPA"), (9.5, "eUICC")]
    for x, label in lanes:
        rect = FancyBboxPatch((x - 0.7, 5.4), 1.4, 0.5, boxstyle="round,pad=0.04,rounding_size=0.1",
                              fc="#333", ec="none")
        ax.add_patch(rect)
        ax.text(x, 5.65, label, ha="center", va="center", color="white", fontsize=10, fontweight="bold")
        ax.plot([x, x], [0.5, 5.4], color="#bbb", ls="--", lw=0.8, zorder=0)

    def msg(y, x_from, x_to, text, note=None):
        ax.annotate("", xy=(x_to, y), xytext=(x_from, y),
                    arrowprops=dict(arrowstyle="->", color="#333", lw=1.4))
        ax.text((x_from + x_to)/2, y + 0.1, text, ha="center", fontsize=9, fontweight="bold")
        if note:
            ax.text((x_from + x_to)/2, y - 0.18, note, ha="center", fontsize=8, color="#555", style="italic")

    msg(4.6, 5.0, 1.5, "POST /gsma/rsp2/esipa/getEimPackage",
        note="eidValue (EID hex)")
    msg(4.0, 1.5, 5.0, "BF51 EuiccPackageRequest (signed by eIM)",
        note="A0 SEQ → 80 eimId, 5A eid, 81 counter, 82 txId, A0 psmoList / A1 ecoList, 5F37 sig")
    msg(3.4, 5.0, 9.5, "POST /api/es10/{eid}/euicc-package")
    ax.text(9.7, 3.0, "eUICC verifies:", fontsize=8.5, color="#a00", fontweight="bold")
    ax.text(9.8, 2.78, "  • eimId matches association", fontsize=8, color="#a00")
    ax.text(9.8, 2.58, "  • counterValue > stored", fontsize=8, color="#a00")
    ax.text(9.8, 2.38, "  • executes PSMO/eCO ops", fontsize=8, color="#a00")
    msg(1.8, 9.5, 5.0, "operation results")
    msg(1.2, 5.0, 1.5, "POST /esipa/provideEimPackageResult",
        note="BF50 → 5A eid + BF51 → A0 → 30 SEQ {80 eimId, 81 ctr, 82 txId, 83 seq, op results}")

    plt.tight_layout()
    plt.savefig(path, bbox_inches="tight"); plt.close(fig)


# Generate all four
fig_trust_hierarchy(os.path.join(IMG_DIR, "trust_hierarchy.png"))
fig_certs_per_actor(os.path.join(IMG_DIR, "certs_per_actor.png"))
fig_es10b_auth_sequence(os.path.join(IMG_DIR, "auth_sequence.png"))
fig_signature_layout(os.path.join(IMG_DIR, "signature_layout.png"))
fig_eim_relay(os.path.join(IMG_DIR, "eim_relay.png"))


# ===========================================================================
# DOCUMENT
# ===========================================================================


def _heading(doc, text, level=1):
    h = doc.add_heading(text, level=level)
    return h


def _para(doc, text, bold=False, italic=False, size=11):
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.bold = bold; r.italic = italic
    r.font.size = Pt(size)
    return p


def _code(doc, text):
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.font.name = "Consolas"
    r.font.size = Pt(9)
    return p


def _shade_cell(cell, color_hex: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), color_hex)
    tc_pr.append(shd)


def _table(doc, header, rows, header_color="305496", widths=None):
    t = doc.add_table(rows=1, cols=len(header))
    t.style = "Light Grid Accent 1"
    hdr_row = t.rows[0]
    for i, h in enumerate(header):
        cell = hdr_row.cells[i]
        cell.text = h
        for r in cell.paragraphs[0].runs:
            r.bold = True
            r.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        _shade_cell(cell, header_color)
    for row in rows:
        rc = t.add_row().cells
        for i, v in enumerate(row):
            rc[i].text = str(v)
    if widths:
        for i, w in enumerate(widths):
            for r in t.rows:
                r.cells[i].width = Inches(w)
    return t


doc = Document()
# Tighter default font
style = doc.styles["Normal"]
style.font.name = "Calibri"
style.font.size = Pt(11)


# ----- TITLE PAGE -----
title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = title.add_run("RSP Certificates and Authentication\n")
r.bold = True; r.font.size = Pt(28); r.font.color.rgb = RGBColor(0x1A, 0x4E, 0xA0)

sub = doc.add_paragraph()
sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = sub.add_run("eIM • SM-DP+ • eUICC — implementation guide for the ConnectX simulator suite")
r.font.size = Pt(13); r.font.color.rgb = RGBColor(0x55, 0x55, 0x55)

doc.add_paragraph()
meta = doc.add_paragraph()
meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
meta.add_run("Specifications: GSMA SGP.22 v3 (RSP) • GSMA SGP.26 v3 (test certificates) • GSMA SGP.32 v1.2 (M2M extensions)").italic = True
doc.add_paragraph()
doc.add_picture(os.path.join(IMG_DIR, "trust_hierarchy.png"), width=Inches(6.5))
doc.add_page_break()


# ----- SECTION 1 — OVERVIEW -----
_heading(doc, "1.  Overview", level=1)
_para(doc,
      "Remote SIM Provisioning (RSP) requires three actors — the eIM (eSIM IoT Manager), the SM-DP+ "
      "(Subscription Manager — Data Preparation +) and the eUICC (the chip in the device) — to mutually "
      "authenticate using ECDSA signatures and X.509 certificate chains. "
      "This document describes which certificates and keys each actor holds, where the trust anchors live, "
      "and how the cryptographic flow plays out across the ESipa, ES9+ and ES10b interfaces.")
_para(doc,
      "Everything in this guide is implemented and verified end-to-end in the ConnectX simulator suite. "
      "Profile downloads, enable/disable/delete operations and ESipa-routed PSMO/eCO commands all complete "
      "with full SGP.22 v3 wire conformance — no test-mode bypasses.")


# ----- SECTION 2 — ACTORS & CERTS -----
_heading(doc, "2.  Certificates per actor", level=1)
doc.add_picture(os.path.join(IMG_DIR, "certs_per_actor.png"), width=Inches(6.7))

_heading(doc, "2.1  SM-DP+ (smdpplus.connectxiot.com)", level=2)
_para(doc, "The SM-DP+ acts as the profile preparation server. It must prove its identity to every eUICC "
           "that downloads a profile, and it does so with end-entity certificates issued under a public test CI.")
_table(doc,
       ["Asset", "Purpose", "File / location", "Algorithm"],
       [
           ["Trust anchor (CI public)",
            "Verifies its own DPauth/DPpb certs were issued under the GSMA Test CI.",
            "ci_nist.crt", "ECDSA P-256"],
           ["CERT.DPauth.SIG + private key",
            "Signs ServerSigned1 in ES9+.InitiateAuthentication (step 3).",
            "smdp_dpauth.crt + .key", "ECDSA P-256"],
           ["CERT.DPpb.SIG + private key",
            "Signs SmdpSigned2 in ES9+.AuthenticateClient (step 5). Binds the BPP to a specific eUICC.",
            "smdp_dpbp.crt + .key", "ECDSA P-256"],
           ["TLS cert",
            "HTTPS only. Independent of RSP signing.",
            "/etc/letsencrypt/...", "Let's Encrypt"],
       ],
       widths=[1.4, 2.8, 1.6, 1.0])

_para(doc, "Trust-anchor details — GSMA SGP.26 RSP2 NIST P-256 Test CI:", bold=True)
_code(doc,
      "Subject:           CN=Test CI, OU=TESTCERT, O=RSPTEST, C=IT\n"
      "Self-signed:       yes (root)\n"
      "SubjectKeyId:      F5:41:72:BD:F9:8A:95:D6:5C:BE:B8:8A:38:A1:C1:1D:80:0A:85:C3\n"
      "Curve / hash:      NIST P-256 (prime256v1) / SHA-256\n"
      "Source:            GSMA SGP.26 v3 test certificate package (osmo-smdpp upstream)\n"
      "Notes:             The PRIVATE key is intentionally not on this box — it lives only with whoever\n"
      "                   issues new DPauth/DPpb certs. We have only the public cert here.")


_heading(doc, "2.2  eUICC simulator (per-EID, in-process)", level=2)
_para(doc, "Each virtual eUICC carries its own three-tier cert chain plus a list of trusted CI roots that it will "
           "accept on the SM-DP+ side. Both sets are independent — the trust list controls verification, the "
           "private key controls signing.")
_table(doc,
       ["Asset", "Purpose", "Source"],
       [
           ["Trusted CI #1 — GSMA Test CI",
            "Lets the eUICC verify any SM-DP+ chain rooted in the public GSMA SGP.26 NIST P-256 CI. Public-only.",
            "Copied to certs/_trusted_cis/gsma_test_ci_nist.crt at sim provisioning"],
           ["Trusted CI #2 — ConnectX self-signed CI",
            "The simulator's own CI (full keypair). Allows offline cert issuance for sim-only chains.",
            "Generated per-deployment at certs/<EID>/ci_*.pem"],
           ["EUM cert (intermediate)",
            "Represents the eUICC manufacturer. Issued by the ConnectX CI.",
            "certs/<EID>/eum_cert.pem"],
           ["eUICC end-entity cert + private key",
            "The per-device cert. Used to sign EuiccSigned1 / EuiccSigned2.",
            "certs/<EID>/euicc_{cert,private}.pem"],
       ],
       widths=[2.0, 3.0, 2.0])

_para(doc, "Key list returned to SM-DP+ in EUICCInfo1.euiccCiPKIdListForVerification:", bold=True)
_code(doc, "[ F54172BDF98A95D65CBEB88A38A1C11D800A85C3,    -- GSMA Test CI (NIST P-256)\n"
           "  CF26CF596C453B64AD88C10F18C473118857FE80 ]   -- ConnectX self-signed CI")
_para(doc, "Both SKIs are advertised. The SM-DP+ picks one (usually the GSMA TestCI, as that is its DPauth root) "
           "via the euiccCiPKIdToBeUsed field of the InitiateAuthentication response.")


_heading(doc, "2.3  eIM (eimserver.connectxiot.com)", level=2)
_para(doc, "Unlike the SM-DP+, the eIM does NOT participate in any X.509 trust chain on the wire. It identifies itself "
           "to the eUICC by an opaque eimId string and a per-association counter (replay protection). The eIM "
           "operator's public key (raw EC) is stored on the eUICC at AddEim time; subsequent requests are signed "
           "with the matching private key.")
_table(doc,
       ["Asset", "Purpose", "File / location"],
       [
           ["eIM signing keypair (raw EC P-256, no X.509)",
            "Signs every EuiccPackageRequest sent over ESipa.",
            "/var/www/eim/storage/app/eim/keys/<eimId>_private.pem"],
           ["eim_registry row (database)",
            "Stores eim_id, eim_id_hmac (for multi-tenant lookup), public-key bytes, counter, supported_protocol.",
            "MySQL table eim_db.eim_registry"],
           ["No CA cert / no GSMA chain",
            "The eIM identity is verified by counter + signature only, not by a cert chain.",
            "—"],
       ],
       widths=[2.4, 2.8, 2.0])


# ----- SECTION 3 — SEQUENCE -----
doc.add_page_break()
_heading(doc, "3.  Authentication flow (profile download)", level=1)
_para(doc, "The 8-step profile-download flow exercises every ECDSA signature, every cert chain validation, "
           "and every key-exchange path in RSP. The sequence below is what our simulator suite executes "
           "verbatim against an SGP.22 v3 SM-DP+. Each labelled phase corresponds to one or more spec "
           "sections (SGP.22 §5.6 for ES9+, §5.7 for ES10b).")
doc.add_picture(os.path.join(IMG_DIR, "auth_sequence.png"), width=Inches(7.0))

_heading(doc, "3.1  Step-by-step explanation", level=2)

steps = [
    ("Step 1 — GetEuiccChallenge (ES10b)",
     "The IPA asks the eUICC for a fresh 16-byte random challenge. This Nonce is the eUICC's "
     "session-binding token; the SM-DP+ must echo it back signed with CERT.DPauth.SIG to prove "
     "freshness."),
    ("Step 2 — GetEuiccInfo1 (ES10b)",
     "The eUICC returns BF20 EUICCInfo1 carrying euiccCiPKIdListForVerification (the SKIs of every "
     "CI root the eUICC trusts) and euiccCiPKIdListForSigning. The SM-DP+ uses these to decide which of "
     "its DPauth certs to use for the signature."),
    ("Step 3 — ES9+.InitiateAuthentication (IPA → SM-DP+)",
     "The IPA forwards the challenge, info1 and the SM-DP+ FQDN. The SM-DP+ responds with: "
     "transactionId, ServerSigned1 (a SEQUENCE with txId/euiccChallenge/serverAddress/serverChallenge), "
     "the 5F37-wrapped serverSignature1 (made with CERT.DPauth.SIG over the canonical DER bytes of "
     "ServerSigned1), the DPauth cert, and the chosen euiccCiPKIdToBeUsed (DER-OCTET-STRING wrapped "
     "as 04 14 <SKI>)."),
    ("Step 4 — ES10b.AuthenticateServer (eUICC verifies SM-DP+)",
     "The IPA strips the 04 14 wrapper from euiccCiPKIdToBeUsed, parses the ServerSigned1 base64 into "
     "fields, strips the 5F37 40 wrapper from the signature, then forwards everything plus the original "
     "ServerSigned1 raw DER bytes to the eUICC. The eUICC checks: (a) ciPKID is in its trust list, "
     "(b) the SM-DP+'s cert chain validates to that CI, (c) ECDSA(serverPubKey, signature, "
     "ServerSigned1Raw) verifies, (d) the euiccChallenge in ServerSigned1 matches the one issued in "
     "step 1. On success the eUICC returns EuiccSigned1 + euiccSignature1 + its own cert + the EUM "
     "cert. EuiccSigned1 carries ctxParams1 echoing the matchingId from the activation code."),
    ("Step 5 — ES9+.AuthenticateClient (IPA → SM-DP+)",
     "The IPA hand-rolls the AuthenticateServerResponse outer wrapper at the byte level: "
     "BF38 → A0 → {EuiccSigned1Raw + 5F37+sig + euiccCert + eumCert}, base64-encodes it, and POSTs "
     "to the SM-DP+. The SM-DP+ verifies the eUICC's cert chain (eUICC → EUM → CI) and the signature "
     "over EuiccSigned1, then looks up a released profile by matchingId. It returns SmdpSigned2 + a "
     "5F37-wrapped smdpSignature2 + the CERT.DPpb.SIG + profile metadata. Critically, "
     "smdpSignature2 is computed over SmdpSigned2 || 5F37 40 || euiccSignature1 — the SM-DP+ binds its "
     "response to the eUICC's prior signature so it can't be replayed against a different session."),
    ("Step 6 — ES10b.PrepareDownload (eUICC verifies SM-DP+ profile-binding)",
     "The eUICC validates the DPpb cert chain (separately from DPauth), then verifies smdpSignature2 "
     "against SmdpSigned2 || 5F37 40 || euiccSignature1 using the DPpb public key. It generates its "
     "one-time public/private keypair (otpk), derives SCP03t session keys via ECDH against the SM-DP+'s "
     "ephemeral key, and returns EuiccSigned2 + a 5F37-wrapped euiccSignature2."),
    ("Step 7 — ES9+.GetBoundProfilePackage",
     "The IPA forwards EuiccSigned2 + signature back to the SM-DP+. The SM-DP+ verifies that euiccSig2 "
     "was made by the same eUICC that authenticated in step 5, then encrypts the actual profile elements "
     "with the derived SCP03t keys and returns the BoundProfilePackage (BF36 envelope)."),
    ("Step 8 — ES10b.LoadBoundProfilePackage",
     "The eUICC walks the BF36 TLV tree, decrypts each profile element with SCP03t (CMAC-verified), "
     "instantiates an ISD-P, installs the profile, and returns a ProfileInstallationResult. Profile is "
     "now stored on the eUICC in the disabled state, ready to be enabled via a separate ESipa command."),
]

for title, body in steps:
    _heading(doc, title, level=3)
    _para(doc, body)


# ----- SECTION 4 — SIGNATURE TBS -----
doc.add_page_break()
_heading(doc, "4.  Signature layouts — what gets signed", level=1)
_para(doc, "Every ECDSA signature in RSP is over canonical DER bytes — exact, byte-for-byte. Any re-encoding "
           "ambiguity (length-form, AUTOMATIC TAGS edge case) breaks the signature. The simulator passes the "
           "original DER through verbatim for verification; never re-encodes between hops.")
doc.add_picture(os.path.join(IMG_DIR, "signature_layout.png"), width=Inches(6.7))

_heading(doc, "4.1  Tag layouts that matter (canonical SGP.22 v3)", level=2)
_table(doc,
       ["Type", "Tag layout", "Notes"],
       [
           ["ServerSigned1",
            "SEQUENCE { [0] txId, [1] euiccChallenge, [3] serverAddress, [4] serverChallenge }",
            "Tags [2] reserved. Re-encoding with [0][1][2][3] produces different DER bytes → sig fails."],
           ["EuiccSigned1",
            "SEQUENCE { [0] txId, [3] serverAddress, [4] serverChallenge, [34] euiccInfo2, ctxParams1 }",
            "ctxParams1 is mandatory in v3, an untagged CHOICE (matchingId carried inside)."],
           ["SmdpSigned2",
            "SEQUENCE { [0] txId, ccRequiredFlag (BOOL), [APPLICATION 73] bppEuiccOtpk OPT }",
            "ccRequiredFlag is untagged BOOLEAN. The OTPK uses [APPLICATION 73] = 5F49 wire tag."],
           ["EuiccSigned2",
            "SEQUENCE { [0] txId, [APPLICATION 73] euiccOtpk, hashCc OPT }",
            "Same APPLICATION-tag pattern; hashCc untagged when present."],
           ["AuthenticateResponseOk",
            "SEQUENCE { EuiccSigned1, [APPLICATION 55] euiccSig1, euiccCert, eumCert }",
            "Outer wrapped in BF38 → A0 (CHOICE alternative 0)."],
       ],
       widths=[1.5, 3.0, 2.6])


# ----- SECTION 5 — eIM RELAY -----
doc.add_page_break()
_heading(doc, "5.  ESipa relay (eIM → eUICC PSMO / eCO commands)", level=1)
_para(doc, "Profile state operations (enable / disable / delete / listProfileInfo, etc.) and eIM configuration "
           "operations (addEim / deleteEim / listEim) are not direct ES10b commands. They flow through ESipa: "
           "the eIM packages them as signed BF51 EuiccPackageRequest messages, the IPA polls and relays them "
           "to the eUICC, and results return to the eIM as BF50 ProvideEimPackageResult.")
doc.add_picture(os.path.join(IMG_DIR, "eim_relay.png"), width=Inches(6.7))

_heading(doc, "5.1  EuiccPackageRequest — what the eIM actually signs", level=2)
_code(doc,
      "BF51 <len>\n"
      "  30 <len>                                    -- EuiccPackageSigned\n"
      "    80 <len> <eimId UTF-8>                    -- the eIM's identifier\n"
      "    5A 10 <eid 16B>                           -- target EID\n"
      "    81 <len> <counterValue>                   -- replay protection\n"
      "    82 10 <transactionId 16B>                 -- session ID (optional)\n"
      "    A0 <len> <psmoList content> | A1 <len> <ecoList content>\n"
      "  5F37 40 <signature 64B>                     -- ECDSA(eIM private key)")

_para(doc, "PSMO action tags (inside A0):", bold=True)
_code(doc,
      "A3 enable      A4 disable      A5 delete      A6 getRAT\n"
      "A7 configureImmediateEnable    A8 setFallbackAttribute    A9 unsetFallbackAttribute\n"
      "BF2D listProfileInfo           BF65 setDefaultDpAddress")

_para(doc, "eCO action tags (inside A1):", bold=True)
_code(doc, "A8 addEim      A9 deleteEim      AA updateEim      AB listEim")

_heading(doc, "5.2  ProvideEimPackageResult — what the eUICC returns", level=2)
_code(doc,
      "BF50 <len>                                   -- ProvideEimPackageResult outer\n"
      "  5A 10 <eid>                                -- REQUIRED for finishQueueItem in eIM\n"
      "  BF51 <len>                                 -- euiccPackageResult\n"
      "    A0 <len>                                 -- euiccPackageResultSigned (constructed)\n"
      "      30 <len>                               -- SEQUENCE\n"
      "        80 <eimId>  81 <counter>  82 <txId>  83 <seqNumber>\n"
      "        <op result tags>")

_para(doc, "Op-result tags read by the eIM:", bold=True)
_code(doc,
      "0x83 / 0xA3 enable result          0x84 / 0xA4 disable        0x85 / 0xA5 delete\n"
      "0xA8 addEim result (constructed)   0x89 / 0xA9 deleteEim\n"
      "0xBF2D listProfileInfo result wraps E3 ProfileInfo entries (5A iccid, 9F70 state, 9F12 name).")


# ----- SECTION 6 — INTEROPERABILITY -----
doc.add_page_break()
_heading(doc, "6.  Interoperability with third-party RSP servers", level=1)
_table(doc,
       ["Direction", "Requirement", "ConnectX status"],
       [
           ["SM-DP+ → eUICC (verify SM-DP+ chain)",
            "eUICC must trust the SM-DP+'s CI root.",
            "✓ eUICC trusts GSMA Test CI as additional root → accepts any SM-DP+ rooted there"],
           ["eUICC → SM-DP+ (verify eUICC chain)",
            "SM-DP+ must trust the eUICC's CI root (real production: GSMA Test CI).",
            "Partial — works against our SM-DP+ today; for strict third parties the eUICC's EUM + eUICC certs need re-issuing under GSMA Test CI (requires its private key from SGP.26 v3 test pack)."],
           ["Sig math byte equivalence",
            "TBS bytes must match canonical SGP.22 v3 layout exactly.",
            "✓ Schema fully aligned (see §4.1)."],
       ],
       widths=[2.2, 2.6, 2.4])

_heading(doc, "6.1  How to upgrade to full third-party interop", level=2)
_para(doc,
      "1.  Obtain GSMA SGP.26 v3 test pack (gsma.com → click-through license). Look for SK_CI_ECDSA_NIST.pem.")
_para(doc,
      "2.  Drop the file into a known location on the deployment box.")
_para(doc,
      "3.  Run the cert regeneration script (~30 min): create a new EUM cert signed by SK_CI_ECDSA_NIST, "
      "then per-EID issue eUICC certs signed by the new EUM. Replace the contents of "
      "certs/<EID>/eum_*.pem and certs/<EID>/euicc_*.pem.")
_para(doc,
      "4.  Sim restart picks up the new chain. The eUICC's euiccCiPKIdListForSigning now reports the GSMA "
      "TestCI SKI; any SGP.22 v3 SM-DP+ that trusts that root will accept our eUICC.")


# ----- SECTION 7 — APPENDIX -----
doc.add_page_break()
_heading(doc, "7.  Appendix — files and tools", level=1)

_table(doc,
       ["File / location", "Purpose"],
       [
           ["smdpplus/scripts/lib/asn1/rsp/rsp.asn",
            "Canonical SGP.22 v3 ASN.1 schema (osmo-smdpp upstream)"],
           ["{euicc,ipa}-simulator/asn1_schemas/rsp_definitions.asn",
            "Simulator-side schema — must match the canonical above"],
           ["euicc-simulator/app/crypto/certificates.py",
            "Per-EID cert chain provisioning + multi-CI trust list"],
           ["euicc-simulator/app/crypto/cert_validator.py",
            "X.509 chain validation against any of the trusted CIs"],
           ["euicc-simulator/app/es10/es10b.py",
            "AuthenticateServer / PrepareDownload signature math"],
           ["ipa-simulator/app/clients/smdp_client.py",
            "ServerSigned1 / SmdpSigned2 wire parsers; hand-rolled AuthenticateServerResponse builder"],
           ["smdpplus/storage/app/certs/",
            "DPauth + DPpb cert + private key + GSMA Test CI public cert"],
           ["euicc-simulator/certs/_trusted_cis/",
            "Additional trust anchors (GSMA Test CI public)"],
           ["euicc-simulator/certs/<EID>/",
            "Per-EID self-signed CI + EUM + eUICC chain"],
       ],
       widths=[3.4, 3.6])

_para(doc, "")
_para(doc, "Verification checklist:", bold=True)
_para(doc, "• Schema byte-equivalence: parse a captured Comprion AuthenticateClient request and diff its "
           "EuiccSigned1 tag tree against ours (same tag layout, different values).")
_para(doc, "• ECDSA signature: feed our SM-DP+'s ServerSigned1 + sig + DPauth cert to "
           "es10b.authenticate_server() and confirm verified=true.")
_para(doc, "• End-to-end: run /transactions/N for any tx with operation = retrieve_eim_package and a "
           "downstream profile_download — every step in §3 should appear with status 200.")


# Save
doc.save(sys.argv[1])
print(f"Wrote {sys.argv[1]}")
